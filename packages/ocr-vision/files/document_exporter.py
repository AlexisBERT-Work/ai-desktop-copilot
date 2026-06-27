"""Generate documents (PDF / DOCX / HTML / Markdown) from Markdown or plain text.

- md/txt : write the content verbatim.
- html   : Markdown -> HTML (python-markdown) with a minimal stylesheet.
- pdf    : Markdown -> HTML -> PDF via xhtml2pdf (pure Python, no system libs).
- docx   : a small block-level Markdown -> python-docx converter.
"""

import logging
import os
import re
from typing import Dict, Any, Optional

log = logging.getLogger(__name__)

SUPPORTED = {"pdf", "docx", "html", "md", "markdown", "txt"}

_EXT_TO_FMT = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".html": "html",
    ".htm": "html",
    ".md": "md",
    ".markdown": "md",
    ".txt": "txt",
}

_CSS = (
    "body{font-family:Helvetica,Arial,sans-serif;font-size:11pt;line-height:1.45;color:#1a1a1a;}"
    "h1{font-size:20pt;}h2{font-size:16pt;}h3{font-size:13pt;}"
    "code,pre{font-family:Courier,monospace;background:#f4f4f4;}"
    "pre{padding:8px;border:1px solid #ddd;}"
    "table{border-collapse:collapse;}td,th{border:1px solid #999;padding:4px 8px;}"
)


def _infer_format(path: str, explicit: Optional[str]) -> str:
    if explicit:
        fmt = explicit.lower()
    else:
        fmt = _EXT_TO_FMT.get(os.path.splitext(path)[1].lower(), "")
    if fmt in ("markdown", "txt"):
        fmt = "md"
    if fmt not in {"pdf", "docx", "html", "md"}:
        raise ValueError(f"Unsupported format '{explicit or os.path.splitext(path)[1]}'. Use pdf, docx, html or md.")
    return fmt


def _md_to_html(content: str, title: Optional[str]) -> str:
    try:
        import markdown as md
    except ImportError:
        raise RuntimeError("markdown not installed. Run: pip install markdown")
    body = md.markdown(content, extensions=["extra", "sane_lists", "tables", "fenced_code", "nl2br"])
    heading = f"<h1>{_escape(title)}</h1>\n" if title else ""
    safe_title = _escape(title or "Document")
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<title>{safe_title}</title><style>{_CSS}</style></head>"
        f"<body>{heading}{body}</body></html>"
    )


def _escape(text: Optional[str]) -> str:
    if not text:
        return ""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _write_text(path: str, text: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)


def _export_pdf(content: str, path: str, title: Optional[str]) -> None:
    try:
        from xhtml2pdf import pisa
    except ImportError:
        raise RuntimeError("xhtml2pdf not installed. Run: pip install xhtml2pdf")
    html = _md_to_html(content, title)
    with open(path, "wb") as f:
        status = pisa.CreatePDF(html, dest=f, encoding="utf-8")
    if status.err:
        raise RuntimeError(f"PDF generation failed ({status.err} error(s))")


# ─── Minimal Markdown -> DOCX ──────────────────────────────────────
_INLINE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`)")


def _add_inline_runs(paragraph, text: str) -> None:
    for part in _INLINE.split(text):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**")) or (part.startswith("__") and part.endswith("__")):
            paragraph.add_run(part[2:-2]).bold = True
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def _export_docx(content: str, path: str, title: Optional[str]) -> None:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    in_code = False
    code_lines: list[str] = []
    for raw in content.splitlines():
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading:
            doc.add_heading(heading.group(2), level=len(heading.group(1)))
            continue

        bullet = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if bullet:
            _add_inline_runs(doc.add_paragraph(style="List Bullet"), bullet.group(1))
            continue

        numbered = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if numbered:
            _add_inline_runs(doc.add_paragraph(style="List Number"), numbered.group(1))
            continue

        _add_inline_runs(doc.add_paragraph(), line)

    if in_code and code_lines:  # unterminated fence — flush what we have
        p = doc.add_paragraph()
        p.add_run("\n".join(code_lines)).font.name = "Courier New"

    doc.save(path)


def export_document(content: str, path: str, fmt: Optional[str] = None, title: Optional[str] = None) -> Dict[str, Any]:
    if content is None:
        raise ValueError("content is required")

    resolved = _infer_format(path, fmt)

    parent = os.path.dirname(os.path.abspath(path))
    os.makedirs(parent, exist_ok=True)

    if resolved == "md":
        _write_text(path, content)
    elif resolved == "html":
        _write_text(path, _md_to_html(content, title))
    elif resolved == "pdf":
        _export_pdf(content, path, title)
    elif resolved == "docx":
        _export_docx(content, path, title)

    size = os.path.getsize(path) if os.path.exists(path) else 0
    return {"path": os.path.abspath(path), "format": resolved, "bytes": size}
