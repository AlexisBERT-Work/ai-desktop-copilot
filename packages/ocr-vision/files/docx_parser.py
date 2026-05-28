"""DOCX file parser using python-docx"""

import logging
from typing import Dict, Any

log = logging.getLogger(__name__)


def parse_docx_file(path: str) -> Dict[str, Any]:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document(path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_text = []

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        tables_text.append("\n".join(rows))

    full_text = "\n\n".join(paragraphs)
    if tables_text:
        full_text += "\n\n[Tables]\n" + "\n\n".join(tables_text)

    return {
        "text": full_text,
        "paragraphCount": len(paragraphs),
        "tableCount": len(doc.tables),
        "metadata": {
            "author": doc.core_properties.author or "",
            "title": doc.core_properties.title or "",
            "created": str(doc.core_properties.created or ""),
        },
    }
